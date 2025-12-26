"""
Document Processor Module
=========================
Handles extraction of text from various document types:
- PDF (native text + OCR for scanned)
- Word documents (.docx)
- Text files (.txt)
- Images (.jpg, .png) via OCR
- CSV files
- SQLite databases

Implements intelligent chunking with overlap for better context retrieval.
"""

import os
import re
import io
import base64
from typing import List, Dict, Any, Optional
import sqlite3

# Document parsing libraries
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


class DocumentProcessor:
    """
    Processes various document types and extracts text content.
    Splits content into overlapping chunks for better RAG performance.
    """
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 100
    ):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Target size of each text chunk (in characters)
            chunk_overlap: Number of overlapping characters between chunks
            min_chunk_size: Minimum chunk size to keep
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
    
    def process_document(self, file_path: str, file_ext: str) -> List[Dict[str, Any]]:
        """
        Main entry point for document processing.
        
        Args:
            file_path: Path to the document file
            file_ext: File extension (e.g., '.pdf', '.docx')
            
        Returns:
            List of chunk dictionaries with text and metadata
        """
        file_ext = file_ext.lower()
        
        # Route to appropriate handler
        handlers = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
            '.csv': self._process_csv,
            '.db': self._process_sqlite
        }
        
        handler = handlers.get(file_ext)
        if not handler:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Extract raw content
        raw_content = handler(file_path)
        
        # Create chunks with metadata
        chunks = self._create_chunks(raw_content)
        
        return chunks
    
    def _process_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Process PDF files - handles both text-based and scanned PDFs.
        Uses PyMuPDF for text extraction, falls back to OCR if needed.
        """
        content = []
        
        if fitz:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc, 1):
                # Try to extract text directly
                text = page.get_text("text")
                
                # If no text found, try OCR
                if not text.strip() and Image and pytesseract:
                    # Render page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    text = pytesseract.image_to_string(img)
                
                if text.strip():
                    content.append({
                        "text": self._clean_text(text),
                        "page": page_num,
                        "source_type": "pdf"
                    })
            
            doc.close()
        
        elif pdfplumber:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text.strip():
                        content.append({
                            "text": self._clean_text(text),
                            "page": page_num,
                            "source_type": "pdf"
                        })
        else:
            raise ImportError("No PDF library available. Install PyMuPDF or pdfplumber.")
        
        return content
    
    def _process_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Process Word documents (.docx)"""
        if not Document:
            raise ImportError("python-docx not installed")
        
        doc = Document(file_path)
        content = []
        
        # Extract paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            full_text.append("\n".join(table_text))
        
        if full_text:
            content.append({
                "text": self._clean_text("\n\n".join(full_text)),
                "page": 1,
                "source_type": "docx"
            })
        
        return content
    
    def _process_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return [{
            "text": self._clean_text(text),
            "page": 1,
            "source_type": "txt"
        }]
    
    def _process_image(self, file_path: str) -> List[Dict[str, Any]]:
        """Process images using OCR"""
        if not Image or not pytesseract:
            raise ImportError("PIL and pytesseract required for image processing")
        
        img = Image.open(file_path)
        
        # Preprocess image for better OCR
        # Convert to grayscale
        if img.mode != 'L':
            img = img.convert('L')
        
        # Perform OCR
        text = pytesseract.image_to_string(img, lang='eng')
        
        if not text.strip():
            return []
        
        return [{
            "text": self._clean_text(text),
            "page": 1,
            "source_type": "image_ocr"
        }]
    
    def _process_csv(self, file_path: str) -> List[Dict[str, Any]]:
        """Process CSV files - converts to readable text format"""
        if not pd:
            raise ImportError("pandas required for CSV processing")
        
        df = pd.read_csv(file_path)
        content = []
        
        # Convert dataframe to readable text
        # Include column headers
        headers = ", ".join(df.columns.tolist())
        text_parts = [f"Columns: {headers}\n"]
        
        # Convert rows to text
        for idx, row in df.iterrows():
            row_text = "; ".join([f"{col}: {val}" for col, val in row.items()])
            text_parts.append(f"Row {idx + 1}: {row_text}")
        
        content.append({
            "text": self._clean_text("\n".join(text_parts)),
            "page": 1,
            "source_type": "csv",
            "row_count": len(df),
            "column_count": len(df.columns)
        })
        
        return content
    
    def _process_sqlite(self, file_path: str) -> List[Dict[str, Any]]:
        """Process SQLite database files"""
        content = []
        
        conn = sqlite3.connect(file_path)
        cursor = conn.cursor()
        
        # Get all tables
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = cursor.fetchall()
        
        for table in tables:
            table_name = table[0]
            
            # Get table schema
            cursor.execute(f"PRAGMA table_info({table_name})")
            columns = [col[1] for col in cursor.fetchall()]
            
            # Get table data (limit for large tables)
            cursor.execute(f"SELECT * FROM {table_name} LIMIT 1000")
            rows = cursor.fetchall()
            
            # Format as text
            text_parts = [f"Table: {table_name}", f"Columns: {', '.join(columns)}", ""]
            
            for row in rows:
                row_text = "; ".join([f"{col}: {val}" for col, val in zip(columns, row)])
                text_parts.append(row_text)
            
            content.append({
                "text": self._clean_text("\n".join(text_parts)),
                "page": 1,
                "source_type": "sqlite",
                "table_name": table_name
            })
        
        conn.close()
        return content
    
    def process_base64_image(self, base64_string: str) -> str:
        """
        Process a base64 encoded image and extract text via OCR.
        
        Args:
            base64_string: Base64 encoded image data
            
        Returns:
            Extracted text from the image
        """
        if not Image or not pytesseract:
            return ""
        
        try:
            # Remove data URL prefix if present
            if "," in base64_string:
                base64_string = base64_string.split(",")[1]
            
            # Decode base64
            image_data = base64.b64decode(base64_string)
            img = Image.open(io.BytesIO(image_data))
            
            # Convert to grayscale for better OCR
            if img.mode != 'L':
                img = img.convert('L')
            
            # Perform OCR
            text = pytesseract.image_to_string(img, lang='eng')
            
            return self._clean_text(text)
        
        except Exception as e:
            print(f"Error processing base64 image: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        # Strip leading/trailing whitespace
        text = text.strip()
        return text
    
    def _create_chunks(self, content: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Split content into overlapping chunks for better retrieval.
        
        Uses a sliding window approach with configurable chunk size and overlap.
        """
        chunks = []
        
        for item in content:
            text = item["text"]
            metadata = {k: v for k, v in item.items() if k != "text"}
            
            # If text is small enough, keep as single chunk
            if len(text) <= self.chunk_size:
                chunks.append({
                    "text": text,
                    **metadata
                })
                continue
            
            # Split into sentences for better chunk boundaries
            sentences = self._split_into_sentences(text)
            
            current_chunk = ""
            current_chunk_sentences = []
            
            for sentence in sentences:
                # Check if adding this sentence exceeds chunk size
                if len(current_chunk) + len(sentence) > self.chunk_size:
                    if current_chunk:
                        chunks.append({
                            "text": current_chunk.strip(),
                            **metadata
                        })
                    
                    # Start new chunk with overlap
                    # Include last few sentences for context
                    overlap_text = ""
                    overlap_sentences = []
                    
                    for s in reversed(current_chunk_sentences):
                        if len(overlap_text) + len(s) > self.chunk_overlap:
                            break
                        overlap_sentences.insert(0, s)
                        overlap_text = s + " " + overlap_text
                    
                    current_chunk = overlap_text + sentence
                    current_chunk_sentences = overlap_sentences + [sentence]
                else:
                    current_chunk += " " + sentence if current_chunk else sentence
                    current_chunk_sentences.append(sentence)
            
            # Add final chunk
            if current_chunk and len(current_chunk) >= self.min_chunk_size:
                chunks.append({
                    "text": current_chunk.strip(),
                    **metadata
                })
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using regex"""
        # Simple sentence splitting - handles common cases
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
